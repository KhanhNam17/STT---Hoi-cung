# components/export_docx.py
#
# Xuất DOCX cho Smart Meeting — tạo TỪ ĐẦU bằng python-docx, KHÔNG cần template.
# Hai hàm chính:
#   - export_to_docx(turns, session_name)             → Biên bản cuộc họp đầy đủ
#   - export_summary_to_docx(summary_text, session_name) → Bản tóm tắt

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Bảng màu pin cho từng người nói (theo thứ tự xuất hiện)
COLORS_RGB = [
    RGBColor(0xE8, 0x52, 0x0A),   # cam (primary brand)
    RGBColor(0x1A, 0x6F, 0xAD),   # xanh dương
    RGBColor(0x2E, 0x8B, 0x2E),   # xanh lá
    RGBColor(0x8B, 0x2E, 0x8B),   # tím
    RGBColor(0x8B, 0x69, 0x14),   # vàng đậm
    RGBColor(0xA0, 0x50, 0x28),   # nâu
]


# ── Helpers ────────────────────────────────────────────────────────────────────
def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run.font.bold = True


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.font.name = "Times New Roman"
    run_l.font.size = Pt(11)
    run_l.font.bold = True
    run_v = p.add_run(value)
    run_v.font.name = "Times New Roman"
    run_v.font.size = Pt(11)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.bold = True


def _add_body_paragraph(doc: Document, text: str, size: int = 12,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _build_color_map(turns) -> dict:
    speakers = list(dict.fromkeys(t.speaker for t in turns))
    return {spk: COLORS_RGB[i % len(COLORS_RGB)] for i, spk in enumerate(speakers)}


def _add_turn_paragraph(doc: Document, turn, color_map: dict) -> None:
    """Một dòng turn: [mm:ss] <Speaker>: <text> (Speaker tô màu)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Mốc thời gian
    ts = f"[{int(turn.start // 60):02d}:{int(turn.start % 60):02d}]  "
    r_ts = p.add_run(ts)
    r_ts.font.name = "Times New Roman"
    r_ts.font.size = Pt(11)
    r_ts.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Tên người nói (đậm + màu)
    r_spk = p.add_run(f"{turn.speaker}: ")
    r_spk.font.name = "Times New Roman"
    r_spk.font.size = Pt(12)
    r_spk.font.bold = True
    r_spk.font.color.rgb = color_map.get(turn.speaker, RGBColor(0, 0, 0))

    # Nội dung
    r_txt = p.add_run(turn.text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(12)


def _new_doc() -> Document:
    """Document trống với lề chuẩn A4."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return doc


# ── 1. Xuất Biên Bản Cuộc Họp đầy đủ (transcript theo từng người nói) ─────────
def export_to_docx(turns, session_name: str = "MEETING") -> bytes:
    """Sinh biên bản cuộc họp DOCX chứa toàn bộ transcript đã phân người nói."""
    doc = _new_doc()
    now = datetime.now()

    _add_title(doc, "BIÊN BẢN CUỘC HỌP")
    _add_meta_line(doc, "Mã cuộc họp", session_name)
    _add_meta_line(doc, "Ngày",        now.strftime("%d/%m/%Y"))
    _add_meta_line(doc, "Thời gian bắt đầu", now.strftime("%H:%M"))
    if turns:
        last_t = max(t.end for t in turns)
        dur = f"{int(last_t // 60)} phút {int(last_t % 60):02d} giây"
        _add_meta_line(doc, "Thời lượng ghi âm", dur)
        speakers = sorted({t.speaker for t in turns})
        _add_meta_line(doc, "Số người nói",   str(len(speakers)))
        _add_meta_line(doc, "Danh sách",      ", ".join(speakers))

    doc.add_paragraph()
    _add_section_heading(doc, "NỘI DUNG CUỘC HỌP")

    color_map = _build_color_map(turns)
    for t in turns:
        _add_turn_paragraph(doc, t, color_map)

    doc.add_paragraph()
    _add_section_heading(doc, "GHI CHÚ")
    _add_body_paragraph(doc, "Biên bản được tạo tự động bằng Smart Meeting "
                              "(STT: Zipformer · Diarization: NeMo Sortformer). "
                              "Nội dung có thể cần biên tập lại trước khi sử dụng chính thức.",
                         size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 2. Xuất Bản Tóm Tắt (text thuần) ──────────────────────────────────────────
def export_summary_to_docx(summary_text: str, session_name: str = "MEETING") -> bytes:
    """Sinh DOCX chứa bản tóm tắt (text thuần, có thể dùng cho NPU Qwen output)."""
    doc = _new_doc()
    now = datetime.now()

    _add_title(doc, "TÓM TẮT CUỘC HỌP")
    _add_meta_line(doc, "Mã cuộc họp", session_name)
    _add_meta_line(doc, "Ngày",        now.strftime("%d/%m/%Y %H:%M"))

    doc.add_paragraph()
    _add_section_heading(doc, "NỘI DUNG TÓM TẮT")

    for line in summary_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # Heading markdown nhẹ: **TIÊU ĐỀ**
        is_heading = line.startswith("**") and line.endswith("**")
        if is_heading:
            run_text = line.strip("*").strip()
            p = doc.add_paragraph()
            r = p.add_run(run_text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.font.bold = True
        else:
            _add_body_paragraph(doc, line, size=12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
